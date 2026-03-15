# app/services/processors/word_processor.py
import docx
from pathlib import Path
from typing import List, Dict, Any
from .pdf_processor import ExtractedContent  # Réutilise la même structure

class WordProcessor:
    """Processeur pour les fichiers Word (.docx)"""
    
    @staticmethod
    def extract_text(file_path: str) -> ExtractedContent:
        """Extrait le texte d'un fichier Word (.docx)"""
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extraction des métadonnées
            metadata = {}
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                }
            
            # Extraction du texte par paragraphe
            full_text = ""
            paragraphs_info = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if para_text:  # Ignorer les paragraphes vides
                    full_text += para_text + "\n\n"
                    
                    paragraphs_info.append({
                        'paragraph_number': i + 1,
                        'text': para_text,
                        'char_count': len(para_text),
                        'word_count': len(para_text.split()),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extraction des tableaux
            tables_text = ""
            tables_info = []
            
            for i, table in enumerate(doc.tables):
                table_text = WordProcessor._extract_table_text(table)
                if table_text:
                    tables_text += f"\n\n[TABLEAU {i+1}]\n{table_text}\n"
                    tables_info.append({
                        'table_number': i + 1,
                        'text': table_text,
                        'rows': len(table.rows),
                        'cols': len(table.columns) if table.rows else 0
                    })
            
            # Combiner tout le texte
            combined_text = full_text + tables_text
            
            # Informations globales
            pages = [{
                'page_number': 1,  # Word n'a pas de concept de page distinct
                'text': combined_text.strip(),
                'char_count': len(combined_text),
                'word_count': len(combined_text.split()) if combined_text else 0,
                'paragraphs': len(paragraphs_info),
                'tables': len(tables_info)
            }]
            
            # Ajouter les infos aux métadonnées
            metadata.update({
                'paragraphs_count': len(paragraphs_info),
                'tables_count': len(tables_info),
                'paragraphs_info': paragraphs_info,
                'tables_info': tables_info
            })
            
            return ExtractedContent(
                text=combined_text.strip(),
                page_count=1,  # Word: 1 "page" logique
                metadata=metadata,
                pages=pages
            )
            
        except Exception as e:
            raise Exception(f"Erreur lors du traitement Word: {str(e)}")
    
    @staticmethod
    def _extract_table_text(table) -> str:
        """Extrait le texte d'un tableau Word"""
        table_text = ""
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_text.append(cell_text)
            
            if any(row_text):  # Si la ligne n'est pas vide
                table_text += " | ".join(row_text) + "\n"
        
        return table_text.strip()